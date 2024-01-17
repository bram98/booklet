#%%
import numpy as np
import matplotlib.pyplot as plt
import os
from pathlib import Path
from glob import glob
from docx import Document
from docx.shared import Inches
import doc2docx
from tqdm import tqdm
import re
import shutil
import panda as pd

from helper_functions import parse_id, init_folder, write_errors

os.chdir(os.path.dirname(__file__))
#%%

data = pd.read_excel('responses.xlsx')
data.set_index('ID', inplace=True)
#%%
def convert_doc_to_docx_windows():
    print('Converting .doc files to .docx files...')
    doc_files = glob('response_data/project_descriptions_processed/*.doc')
    for doc_file in tqdm(doc_files):
        doc2docx.convert(doc_file)
    print('\nDone converting .doc files to .docx files')
    
    print('Removing old .doc files...')
    for doc_file in tqdm(doc_files):
        Path(doc_file).unlink( missing_ok=True)
    
    print('\nDone removing old .doc files...')

def inspect_file_types():
    filenames = glob('response_data/project_descriptions_renamed/*')
    for filename in filenames:
        ext  = Path(filename).suffix
        if ext != '.docx':
            print(Path(filename).name)

def write_file_type_errors():
    filenames = glob('response_data/project_descriptions_processed/*')
    filename_ids = []
    filename_errs = []
    for filename in filenames:
        file_path = Path(filename)
        ext = file_path.suffix
        if ext != '.docx':
            filename_ids.append( int( parse_id( file_path.name ) ) )
            filename_errs.append( (f'The file type of your project description '
                                   f'is not supported. Your file type: {ext}. '
                                   f'Supported: .doc or .docx.'
                                   ))
    print(f'Detected {len(filename_errs)} wrong projection description file types.')
    write_errors('proj_description', filename_ids, filename_errs)
    # write_errors('figure', caption_ids, caption_errs)
# convert_doc_to_docx_windows()
# inspect_file_types()

#%%

# from glob import escape
filenames = glob(r'response_data/project_descriptions_renamed/*')
target_dir = Path(r'./response_data/project_descriptions_processed')
MODIFY_FILES = True
if MODIFY_FILES:
    init_folder('project_descriptions_processed')
    for filename in filenames:
        src_path = Path(filename)
        target_path = target_dir/src_path.name
        shutil.copy(src_path, target_path)
        
    convert_doc_to_docx_windows()
    write_file_type_errors()
#%%
filenames = glob('response_data/project_descriptions_renamed/*.docx')


# reference_regex = re.compile(r'(^\[\d\])|references:?', flags=re.M|re.I)
reference_regex = re.compile(r'(\[.*?\d.*?\])', flags=re.M|re.I)
info_regex = re.compile(r'^(sponsors?)|(supervisors?)|(keywords)', flags=re.M|re.I)
figure_regex = re.compile(r'^fig(ure)?', flags=re.I)
email_regex = re.compile(r'@', flags=re.I)
# lengths = []
ref_list = []
i=0
for filename in filenames[:]:
    doc = Document(filename)
    # print( [len(p.text) for p in doc.paragraphs] )
    paras = [p.text for p in doc.paragraphs]
    # print(i)
    i+=1
    for para in paras:
        if len(para) == 0:
            continue
        if not reference_regex.search(para) is None:
            # print(para)
            # print(reference_regex.search(para)[0])
            # print(Path(filename).name)
            ref_list.append(get_id(filename))
            # print('')
            # print(len(para))
            break
        # if not info_regex.search(para) is None:
        #     pass
        
        # if not figure_regex.search(para) is None:
        #     # print(para)
        #     pass
        
        # if not email_regex.search(para) is None:
        #     # print(para)
        #     pass
        #     # break
        
        # if len(para) < 250:
        #     # print(para)
        #     print('')
    # lengths.append([len(p.text) for p in doc.paragraphs])
    
ref_list = sorted(ref_list)
print(ref_list)
# lengths_flattened = [length for row in lengths for length in row]

#%%
wrong_proj_descriptions = [
    'Claudia Keijzer',
    'Thimo Jacobs',
    'Just Pe',
    'Huygen Jobsis',
    'Oscar Brandt',
    'Roos Grote',
    'Komal',
    'Matt Peerlings',
    'Pim Witte',
    'Laurens Mandemaker',
    'Pauline Julika',
    'Suzan Schoemaker',
    'Marnix Vreugdenhil',
    'Jim de Ruiter',
    'Maaike Vink',
    'Diogo Vieira',
    'Sofie Ferwerda',
    'George Tierney',
    'Rodolfo Subert',
    'Sepideh',
    'Bettinea',
    'Joyce Kromwijk',
    'Jesse Buckmann',
    'Jelle Kranenborg',
    'Albaraa',
    'Hanya',
    'Kelly Brouwer',
    'Rian Ligthart',
    'Yuang Piao',
    'Roel Biene',
    'Xinwei Ye',
    'Robin Vogel',
    'Mengwei Li',
    'Jesse Steenhoff',
    'Maartje Otten',
    'Qijun Che',
    'Nicolette Maaskant',
    'Jan Cuperus',
    'Riccardo Reho',
    'Luuk Stringer',
    'Zahra',
    'Errikos',
    'Gerardo',
    'Edwin Armando',
    'Sebastian Rejman',
    'Sibylle',
    'Sander Vonk',
    'Luke Riddel',
    'Claire Seitzinger',
    'Monica Conte',
    'Thomas Resz',
    'Thijs ter Rele',
    'Ruizhi Yang',
    'Hui Wang',
    'Xiang Yu',
    'Auke Vlasblom',
    'Jiaorang Yan',
    'Chunning Sun',
    'Zixiong Wei',
    'Montserrat',
    'Karan',
    'Disha',
    'David Villaron',
    'Weizhe Zhang'
    ]
vowel_regex = re.compile(r'[aeiou]')  # To match á î è etc.
email_adresses = ''
# for name in wrong_proj_descriptions:
#     name_regex = re.compile(re.sub(vowel_regex, '.', name))
#     if sum(data['Full Name'].str.contains(name_regex)) != 1:
#         print(name)
#     person = data.loc[data['Full Name'].str.contains(name_regex)]
#     email_adress = person['Email'].item()
#     print(person['Project Type'].item())
#     email_adresses += f"{email_adress},"

people2 = []
for(ID, person) in data.iterrows():
    # print(person['Project Type'] is np.nan)
    if person['Project Type'] is np.nan:
        print(person['Full Name'])
