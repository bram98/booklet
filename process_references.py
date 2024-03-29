#%%
import numpy as np
import pandas as pd
# import PIL
from  glob import glob
import os
from os.path import splitext
# import shutil
from unidecode import unidecode
from urllib.parse import urlparse, parse_qs, unquote
import re
from tqdm import tqdm
from pandas.api.types import is_integer_dtype
from pandas.api.types import is_any_real_numeric_dtype
from pathlib import Path
from docx import Document

os.chdir(os.path.dirname(__file__))
from helper_functions import init_folder, parse_id, copy_file, fuzzy_equal, replace_accents, valid_citation_sequence, write_errors
import bibtexparser_v3
from bibtexparser_v3 import make_bibliography

#%%
# Read in the data as a pandas dataframe
data = pd.read_excel('responses.xlsx')
data.set_index('ID', inplace=True)
#%%

def write_errors_references(person_id, excel_file):
    id_list = []
    err_list = []
    df = pd.read_excel(excel_file)

    if df.columns.size != 6:
        id_list.append(person_id)
        err_list.append('Please do not modify the number of columns in the references excel sheet.')
        return  # Cannot process any further if the number of columns isn't even correct
    if not is_integer_dtype(df.iloc[:,0]):
        id_list.append(person_id)
        err_list.append('Reference error. The first column with reference numbers does not contain (only) numbers.')
    if not is_integer_dtype(df.iloc[:,3]):
        id_list.append(person_id)
        err_list.append('Reference error. The fourth column with journal volume/issue does not contain (only) numbers.')
    if not is_integer_dtype(df.iloc[:,5]):
        id_list.append(person_id)
        err_list.append('Reference error. The sixth column with publication year does not contain (only) numbers.')
    for i in range(df.columns.size):
        if np.any(df.iloc[:,i].isna()):
            id_list.append(person_id)
            err_list.append(f'Reference error. Column {i+1} contains some empty cells')

    write_errors('references', id_list, err_list)



def check_page_range(pr):
    if is_any_real_numeric_dtype(type(pr)) and np.isnan(pr):
        return False
    else:
        return True
    
def escape_markdown(match: re.Match):
    return f'\&#{ord(match.group())};'

def replace_accents_markdown(string: str):
    # Don't replace the following accents; these will break the Bib parsing
    non_alphanumeric = re.compile(r'([^a-zA-Z0-9\@\{\},=\n\ \t\-])', re.UNICODE)
    string = re.sub( non_alphanumeric, escape_markdown, string )
    return string

def docxtobib(docx_file: str):
    docx_path = Path(docx_file)
    doc = Document(docx_file)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    fullText = '\n'.join(fullText)
    # fullText2 = unidecode(fullText)
    # fullText2 = fullText2.replace('"','')
    fullText2 = replace_accents_markdown(fullText)
    # fullText2 = fullText
    if fullText2 != fullText:
        # print('Found unicode! \n', fullText)
        print('Found unicode! \n', Path(docx_file).stem)
    with open(docx_path.with_suffix('.bib'), 'w') as file:
        file.write(fullText2)

def xlsx2bib(xlsx_file: str):
    global df_, s, s2
    # Replace file extension
    base_name = splitext(xlsx_file)[0]
    bib_file = base_name + '.bib'
    
    # Parse excel file
    df = pd.read_excel(xlsx_file)
    df = df.rename({'number':'Number'}, axis='columns')
    df = df.sort_values(by='Number')
    def try_convert(row, column_name, type_):
        try:
            result = type_(row.filter(regex=re.compile(column_name, flags=re.I)).iloc[0])
            if result == 'nan':
                raise
            return result
        except:
            return ''
        
    emdash= u'\u2014'
    
    
    et_al_regex = re.compile('(?<!and)\s*?(\,?\s?et al\.?)')

    # For each row in the xlsx file, generated an entry in the bib file
    with open(bib_file, 'w') as fh:
        for i, row in df.iterrows():
            try:
                author = replace_accents_markdown(row['First author'].replace('&', chr(92)+'&').strip())
                author = re.sub(et_al_regex, r' and \1', author)  # convert 'et al'  into separate author, which will then be removed and replaced by et al because only 1 author is removed.
                journal = replace_accents_markdown(try_convert(row, 'Journal', str).strip())
                issue = try_convert(row, 'Issue', int)
                pages = replace_accents_markdown( try_convert(row, 'page range', str).replace(emdash, '-').replace('-', '--').strip() )
                if re.compile('\d').search(pages):
                    pages_str = f'pages={{{pages}}},\n'
                else:
                    pages_str = ''
                    print(pages, try_convert(row, 'page range', str))#, row.filter(regex=re.compile('pages', flags=re.I)).iloc[0])
                
                fh.write(f"""
@article{{ref{row['Number']},
    author={{{author}}},
    journal={{{journal}}},
    number={{{issue}}},
    {pages_str}
    year={{{try_convert(row, 'year', str)}}},
}}
                """)
            except Exception as e:
                my_row = row
                print(base_name, row, repr(e))

def cite_replace_sb(match: re.Match):
    '''
    Replaces \cite{somecite} in the text by [1] and keeps track in cite_dict. Helper function to be used in regex.
    Make sure you set cite_counter to zero and cite_dict to {} before using
    '''
    global cite_counter, cite_dict   
    
    citation_names = match.groups(0)[0]
    citation_names = citation_names.split(',')
    citation_names = [c.strip() for c in citation_names]
    result = ''
    for citation_name in citation_names:
        if not citation_name in cite_dict:
            cite_counter += 1   
            cite_dict[citation_name] = cite_counter
            # return f'[{cite_counter}]'
        result += ',' + str(cite_dict[citation_name])
    result = result[1:]
    print(result)
    return f'[{result}]'

def cite_replace_ref(match: re.Match):
    global cite_counter
    return f'ref{cite_counter}'

def replace_cite_names_by_number(bib_file: str, text: str) -> tuple[str, list[str]]:
    '''
    Replaces all occurences of \cite{somecite} by [1] in text. Also replace occurences
    of somecite by ref1 in corresponding bibfile.
    '''
    global cite_counter, cite_cite_regex, cite_dict
    
    cite_counter = 0
    cite_dict = {}
    text_replaced = cite_cite_regex.sub(cite_replace_sb, text)
    with open(bib_file, 'r') as file:
        content = file.read()
    new_citations = []
    
    cite_counter = 0
    for (citation, citation_number) in cite_dict.items():
        cite_counter = citation_number
        print(citation)
        (content, hits) = re.subn(citation, cite_replace_ref, content)
        new_citations.append(f'ref{cite_counter}')
        if hits != 1:
            raise Exception(f'Could not find reference \'{citation}\' in file {bib_file}')
    
    with open(bib_file, 'w') as file:
        file.write(content)
    return (text_replaced, new_citations)

#%%
src_folder = './response_data/references_renamed/'
dest_folder = './response_data/references_processed/'
reference_paths = glob(src_folder + '*')

MODIFY_FILES = True

if MODIFY_FILES:
    init_folder('references_processed')
    
'''
Convert excel and docx to .bib
'''
for reference_path in reference_paths[:]:
    ref_path = Path(reference_path)
    ref_name = ref_path.stem
    # if not 'Bram' in ref_name:
    #     continue
    ref_path = copy_file(
        src_folder='references_renamed', 
        dest_folder='references_processed',
        old_filename=ref_path.name,
        new_filename='references',
        name=ref_name[11:]
        )
    ref_path = Path(ref_path)
    print(ref_path)
    try:
        if ref_path.suffix == '.docx':
            docxtobib( ref_path.absolute() )
        elif ref_path.suffix == '.xlsx':
            xlsx2bib(ref_path.absolute())
        elif ref_path.suffix == '.bib' or ref_path.suffix == '.md':
            pass
    except Exception as e:
        print(f'[{ref_name[11:]}] {repr(e)}')
        
#%%
# src_folder = './response_data/references_renamed/'
dest_folder = './response_data/references_processed/'
bib_paths = glob(dest_folder + '*.bib')
proj_description_paths = glob('./response_data/project_descriptions_processed/*')


cite_sb_regex = re.compile(r'(\[.*?\d.*?\])') # Matches [1] and [ 1 ]
# cite_sb_regex2 = re.compile(r'\[.*?(\d+)\D*?(\d+)?.*?\]') # Matches [1], [1, 2] and [8-10]
cite_sb_regex2 = re.compile(r'\[.*?(\d+)(?:(?:\D{1,3})?(\d+))*.*?\]') # Matches [1], [1, 2, 3] and [8-10]. Captures first and last number.
cite_cite_regex = re.compile(r'\\cite\{(.+?)\}')

cite_compound_period = re.compile(r'')
for bib_path in bib_paths:
    bib_path = Path(bib_path)
    person = bib_path.stem[11:]

    proj_descriptions = list(filter(lambda file: person in file, proj_description_paths ))
    
    if len(proj_descriptions) != 1:
        # print(proj_descriptions)
        raise Exception(f'[{person}] Found {len(proj_descriptions)} project descriptions!')
        
    proj_description = proj_descriptions[0]
    
    doc = Document(proj_description)
    # print( [len(p.text) for p in doc.paragraphs] )
    paras = '###'.join( [p.text for p in doc.paragraphs] )

    # Find all citations of the type [1]
    numbers = cite_sb_regex2.findall(paras)
    numbers = np.array(numbers, dtype=str).flatten()
    numbers = np.array(list(filter(lambda x: x, numbers)), dtype=int)
    if len(numbers) == 0:
        cite_cites = cite_cite_regex.findall(paras)
        print(cite_cites, person)
        if len(cite_cites) > 0:
            citations = [citation for citation_list in cite_cites for citation in citation_list.split(',') ]
            
            # Replace \cite{somecite} by [1] in text and by ref1 in .bib file
        
            (paras, citations) = replace_cite_names_by_number(bib_path, paras)
        else:
            citations = ['*']
    else:
        # Found numbers in text. Use largest number 
        try:
            assert valid_citation_sequence(numbers)
        except:
            print(f'Error: citations not a valid sequence{person}: {numbers}')
            print(paras)
        largest_ref = np.max(numbers)
        citations = [f'ref{n}' for n in range(1, largest_ref + 1)]
    # print(citations)
    citations = [citation.strip() for citation in citations]
    
    
    cite_sb_period_regex = re.compile(r'\ ?(\[.*?\d+\D*?(\d+)?\])\.')
    period_space_cite_regex = re.compile(r'\.\ (\[.*?\d+\D*?(\d+)?\])')
    # print()
    # print(paras)
    # if 'Suzan' in person:
    #     print(citations, 'suz')
    #     print(paras)
    paras = re.sub(cite_sb_period_regex, r'.\1', paras)
    paras = re.sub(period_space_cite_regex, r'.\1', paras)
    
    try:
        make_bibliography(bib_path, dest_folder + bib_path.stem, citations=citations)
        pass
        # with open(dest_folder + bib_path.stem, 'rw') as file:
        #     file.write(file.read())
        
    except Exception as e:
        print(person, bib_path, repr(e))
        # raise e
    try:
        paras_list = paras.split('###')
        for index, para in enumerate(doc.paragraphs):
            para.text = paras_list[index]
        doc.save(proj_description)
    except Exception as e:
        print(person)
        raise e



